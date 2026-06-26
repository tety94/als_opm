import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def save_table_to_word(
    title: str,
    caption: str,
    headers: list,
    rows: list,
    filename: str,
    footnotes: list = None,
    output_dir: str = "output/words",
):
    """
    Create a Word document containing a formatted table.

    Parameters
    ----------
    title : str
        Title placed above the table.
    caption : str
        Caption placed below the table.
    headers : list[str]
        List of column headers.
    rows : list[list]
        List of table rows (each row is a list of cell values).
    filename : str
        Output filename (.docx).
    footnotes : list[str], optional
        List of footnotes to append at the end.
    output_dir : str
        Directory where the file will be saved.
    """

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Create document
    doc = Document()

    # Title
    title_par = doc.add_paragraph()
    run = title_par.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    # Add some spacing
    doc.add_paragraph("")

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    # Data rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Caption
    doc.add_paragraph("")
    cap_par = doc.add_paragraph()
    cap_run = cap_par.add_run(caption)
    cap_run.italic = True

    # Footnotes
    if footnotes:
        doc.add_paragraph("")
        for fn in footnotes:
            fn_par = doc.add_paragraph(style="List Number")
            fn_par.add_run(fn)

    # Save file
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path
